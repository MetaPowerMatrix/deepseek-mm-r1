#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
将标签图片排版成Word文档
功能：
- 将tags目录下的所有图片添加到Word文档中
- 设置页面为A4纸张大小
- 页边距为0.5cm
- 每页按2列6行排列图片
"""

import os
import logging
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# 设置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

def create_print_document(images_dir, output_file, cols=2, rows=6):
    """
    创建一个用于打印的Word文档，包含排列好的标签图片
    
    参数:
        images_dir: 包含图片的目录路径
        output_file: 输出的Word文档路径
        cols: 每页的列数（默认2列）
        rows: 每页的行数（默认6行）
    """
    try:
        # 获取所有图片文件并排序
        image_files = sorted([
            f for f in os.listdir(images_dir) 
            if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))
        ])
        
        if not image_files:
            logging.error(f"在 {images_dir} 目录中没有找到图片文件")
            return False
        
        logging.info(f"找到 {len(image_files)} 个图片文件")
        
        # 创建Word文档
        doc = Document()
        
        # 计算A4纸张尺寸（毫米）
        # A4: 210mm x 297mm
        page_width = 210
        page_height = 297
        
        # 设置页边距（左右0.5cm = 5mm，上下0mm）
        margin_lr = 5  # 左右边距
        margin_tb = 0  # 上下边距
        
        # 可用区域尺寸
        usable_width = page_width - 2 * margin_lr
        usable_height = page_height - 2 * margin_tb
        
        # 计算每个单元格的尺寸
        cell_width = usable_width / cols
        cell_height = usable_height / rows
        
        # 设置页面属性（A4纸张、左右页边距5mm，上下页边距0mm）
        section = doc.sections[0]
        section.page_width = Mm(page_width)
        section.page_height = Mm(page_height)
        section.left_margin = Mm(margin_lr)
        section.right_margin = Mm(margin_lr)
        section.top_margin = Mm(margin_tb)
        section.bottom_margin = Mm(margin_tb)
        section.orientation = WD_ORIENT.PORTRAIT  # 纵向
        
        # 计算每页可以放多少张图片
        images_per_page = cols * rows
        
        # 计算需要多少页
        total_pages = (len(image_files) + images_per_page - 1) // images_per_page
        logging.info(f"总共需要 {total_pages} 页")
        
        # 按页处理图片
        for page in range(total_pages):
            # 在新页前添加分页符（第一页除外）
            if page > 0:
                doc.add_page_break()
            
            # 创建表格用于排列图片
            table = doc.add_table(rows=rows, cols=cols)
            table.autofit = False
            
            # 调整表格样式：无边框、精确宽度
            table.style = 'Table Grid'
            # 移除表格边框
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._element.tcPr
                    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                         f'<w:top w:val="nil"/>'
                                         f'<w:left w:val="nil"/>'
                                         f'<w:bottom w:val="nil"/>'
                                         f'<w:right w:val="nil"/>'
                                         f'</w:tcBorders>')
                    tcPr.append(tcBorders)
            
            # 设置表格单元格宽度和高度
            for row_idx in range(rows):
                row = table.rows[row_idx]
                row.height = Mm(cell_height)
                
                for col_idx in range(cols):
                    cell = row.cells[col_idx]
                    cell.width = Mm(cell_width)
                    
                    # 计算当前图片的索引
                    img_idx = page * images_per_page + row_idx * cols + col_idx
                    
                    # 如果图片索引有效，添加图片
                    if img_idx < len(image_files):
                        img_path = os.path.join(images_dir, image_files[img_idx])
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        try:
                            # 不指定宽度和高度，保持图片原始尺寸
                            run = paragraph.add_run()
                            run.add_picture(img_path)
                            logging.info(f"添加图片: {img_path}")
                        except Exception as e:
                            logging.error(f"添加图片 {img_path} 时出错: {str(e)}")
                    
        # 确保输出目录存在
        output_dir = Path(output_file).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 保存文档
        doc.save(output_file)
        logging.info(f"文档已保存到: {output_file}")
        
        return True
        
    except Exception as e:
        logging.error(f"创建打印文档时出错: {str(e)}")
        import traceback
        logging.error(f"错误堆栈: {traceback.format_exc()}")
        return False

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(description='将标签图片排版到Word文档中')
    parser.add_argument('--input-dir', '-i', default='tags',
                       help='输入图片目录，默认为tags')
    parser.add_argument('--output-file', '-o', default='标签打印文档.docx',
                       help='输出Word文档路径，默认为标签打印文档.docx')
    parser.add_argument('--cols', '-c', type=int, default=2,
                       help='每页列数，默认为2')
    parser.add_argument('--rows', '-r', type=int, default=6,
                       help='每页行数，默认为6')
    
    # 解析命令行参数
    args = parser.parse_args()
    
    # 创建打印文档
    success = create_print_document(
        images_dir=args.input_dir,
        output_file=args.output_file,
        cols=args.cols,
        rows=args.rows
    )
    
    if success:
        logging.info(f"打印文档创建成功: {args.output_file}")
        print(f"文档已创建: {args.output_file}")
    else:
        logging.error("打印文档创建失败")
        print("文档创建失败，查看日志了解详情")

if __name__ == "__main__":
    main() 